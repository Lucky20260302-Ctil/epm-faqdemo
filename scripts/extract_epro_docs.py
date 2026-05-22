#!/usr/bin/env python3
"""Extract HKTDC EPRO functional requirements from DOCX and generate Obsidian notes."""
import os, re
from docx import Document

BASE = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'FAQ_test')
SRC = r'D:\ObsidianDB\TDC\Attachments'

# Document mapping: (filename, category_dir, category_label, module_tag)
DOCS = [
    ('HKTDC Functional Requirement Specification - General Requirements.docx',
     '08_EPRO_System', 'EPRO 系統基礎', 'general'),
    ('HKTDC Functional Requirement Specification - RFQ and Tender (Pre-Tender Stage).DOC',
     '09_EPRO_PreTender', 'EPRO 招標前階段', 'pre-tender'),
    ('HKTDC Functional Requirement Specification - RFQ and Tender (Tender Stage).DOC',
     '10_EPRO_TenderStage', 'EPRO 招標階段', 'tender'),
    ('HKTDC Functional Requirement Specification - RFQ and Tender (Post-Tender Stage).DOC',
     '11_EPRO_PostTender', 'EPRO 招標後階段', 'post-tender'),
    ('HKTDC Functional Requirement Specification - Supplier Management.docx',
     '12_EPRO_Supplier', 'EPRO 供應商管理', 'supplier'),
    ('HKTDC Functional Requirement Specification - Report Requirements.docx',
     '13_EPRO_Reports', 'EPRO 報表', 'reports'),
    ('HKTDC Functional Requirement Specification - RFQ and Tender (Others).docx',
     '14_EPRO_Others', 'EPRO 其他', 'others'),
]

def extract_content(doc_path):
    """Extract functional requirements from a DOCX file."""
    doc = Document(doc_path)

    requirements = []
    use_cases = []
    current_req = None
    current_section = None
    in_use_cases = False

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name

        if not text:
            continue

        # Detect section boundaries
        if 'USE CASE' in text.upper() and style.startswith('Heading'):
            in_use_cases = True
            continue
        if 'FUNCTIONAL REQUIREMENT' in text.upper() and style.startswith('Heading'):
            in_use_cases = False
            current_section = 'requirements'
            continue

        # Extract H3 requirement headers (FR-XXX-XXX format)
        if style.startswith('Heading 3') and not in_use_cases:
            req_id_match = re.search(r'(FR-[A-Z]+-\d+)', text)
            if req_id_match:
                current_req = {
                    'id': req_id_match.group(1),
                    'title': text,
                    'description': [],
                    'tables': []
                }
                requirements.append(current_req)
            elif current_req:
                current_req['description'].append(text)
        elif style.startswith('Heading 3') and in_use_cases:
            uc_match = re.search(r'(RTR-\d+|SMR-\d+|GR-\d+)', text)
            if uc_match:
                uc = {
                    'id': uc_match.group(1),
                    'title': text,
                    'description': []
                }
                use_cases.append(uc)
                current_req = uc
        elif current_req and not style.startswith('Heading'):
            current_req['description'].append(text)

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            # Try to associate table with nearest requirement
            reqs = requirements if not in_use_cases else use_cases
            if reqs:
                reqs[-1].setdefault('tables', []).append(rows)

    return requirements, use_cases

def sanitize_filename(text):
    """Create a safe filename from text."""
    text = re.sub(r'[<>:"/\\|?*]', '', text)
    text = text.replace(' ', '-').replace('--', '-')
    return text[:80]

def generate_note(category_dir, entry, entry_type='requirement'):
    """Generate a single Obsidian note."""
    req_id = entry['id']
    title = entry['title'].replace('\t', ' ').strip()
    desc = '\n'.join(entry.get('description', []))

    # Build note content
    parts = [
        '---',
        f'title: "{req_id}: {title[:100]}"',
        f'tags: [epro, {category_dir.split("_",1)[1] if "_" in category_dir else category_dir}]',
        '---',
        '',
        f'# {req_id}: {title}',
        '',
        '## 需求描述',
        '',
        desc if desc else '_無詳細描述_',
    ]

    # Add tables if present
    if entry.get('tables'):
        parts.append('')
        parts.append('## 相關資料表')
        parts.append('')
        for ti, table in enumerate(entry['tables']):
            if len(table) <= 1:
                continue
            parts.append(f'| {" | ".join(table[0])} |')
            parts.append(f'| {" | ".join(["---"] * len(table[0]))} |')
            for row in table[1:]:
                # Pad row to match header length
                padded = row + [''] * (len(table[0]) - len(row))
                parts.append(f'| {" | ".join(padded[:len(table[0])])} |')
            parts.append('')

    filename = sanitize_filename(f'{req_id}-{title[:60]}')
    filepath = os.path.join(BASE, category_dir, f'{filename}.md')
    os.makedirs(os.path.dirname(filepath), exist_ok=True)

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write('\n'.join(parts))

    return filepath

def main():
    total_reqs = 0
    total_ucs = 0

    for doc_filename, cat_dir, cat_label, mod_tag in DOCS:
        doc_path = os.path.join(SRC, doc_filename)
        if not os.path.exists(doc_path):
            print(f'SKIP (not found): {doc_filename}')
            continue

        print(f'Processing: {doc_filename[:60]}...')
        reqs, ucs = extract_content(doc_path)

        for req in reqs:
            generate_note(cat_dir, req, 'requirement')

        for uc in ucs:
            generate_note(cat_dir, uc, 'usecase')

        print(f'  Requirements: {len(reqs)}, Use Cases: {len(ucs)}')
        total_reqs += len(reqs)
        total_ucs += len(ucs)

    print(f'\nTotal: {total_reqs} requirements + {total_ucs} use cases = {total_reqs + total_ucs} notes')
    print('Done.')

if __name__ == '__main__':
    main()
